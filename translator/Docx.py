__author__ = 'dor'


from docx import Document
from docx.shared import Inches
from .models import Paper,Paragraph,Translated_Paragraph
import RegexTxt
from django.utils.encoding import smart_unicode
from HTMLParser import HTMLParser
from docx import Document
from django.utils.encoding import smart_unicode
# create a subclass and override the handler methods
from docx.text.run import Font, Run

class MyHTMLParser(HTMLParser):
    def handle_starttag(self, tag, attrs):
        if tag == 'p':
            self.p = document.add_paragraph()
        elif tag == 'br' and self.p is not None:
            run=self.p.add_run()
            run.add_break()
            # run.breaks

        elif tag == 'ol':
            self.numbers=True
        elif tag =='ul':
            self.numbers=False
        elif tag == 'strong':
            self.bold=True
        elif tag == 'em' or tag =='i':
            self.i=True

    def handle_endtag(self, tag):
        if tag == 'p':
            self.p = None
        elif tag == 'strong':
            self.bold=False
        elif tag == 'em' or tag =='i':
            self.i=False

    def handle_data(self, data):
        if self.lasttag == 'p' and self.p is not None :
            self.p.add_run(data)
        elif self.lasttag == 'strong' and self.p is not None:
            add_run=self.p.add_run(data)
            add_run.bold=self.bold
            # add_run.font.rtl=True
        elif (self.lasttag == 'i' or self.lasttag == 'em') and self.p is not None :
            add_run = self.p.add_run(data)
            add_run.italic = self.i
            # add_run.font.rtl=True
        elif self.lasttag == 'li' :
            if self.numbers:
                document.add_paragraph(data, style='ListNumber')
            else:
                document.add_paragraph(data, style='ListBullet')
        elif self.lasttag == "title":
            document.add_heading(data,0)
        elif self.lasttag == 'h1':
            document.add_heading(data,level=1)
        elif self.lasttag == 'h2':
            document.add_heading(data,level=2)
        elif self.lasttag == 'h3' or self.lasttag == 'h4' or self.lasttag == 'h5':
            document.add_heading(data,level=3)


parser = MyHTMLParser()
document = Document()

def WriteDocx(paper):

    document.add_heading(paper.name, 0)
    htmlPaper=""
    Paras=Paragraph.objects.filter(paperId=paper)
    for para in Paras:
        if(Translated_Paragraph.objects.filter(paraId=para).exists()):
            htmlPaper+=Translated_Paragraph.objects.filter(paraId=para)[0].txt

        else:
            htmlPaper+="<p>"+para.txt+"</p>"
    parser.feed(smart_unicode(htmlPaper))
    return document

def original_file(paper):

    document.add_heading(paper.name, 0)
    htmlPaper=""
    Paras=Paragraph.objects.filter(paperId=paper)
    for para in Paras:
            htmlPaper+="<p>"+para.txt+"</p>"
    parser.feed(smart_unicode(htmlPaper))
    return document

def original_file_bold(paper,num):
    document.add_heading(paper.name, 0)
    htmlPaper=""
    Paras=Paragraph.objects.filter(paperId=paper)
    for para in Paras:
        if para.num == int(num):
            htmlPaper += "<p><strong>"+para.txt+"</strong></p>"
            continue
        htmlPaper += "<p>"+para.txt+"</p>"
    parser.feed(smart_unicode(htmlPaper))
    return document